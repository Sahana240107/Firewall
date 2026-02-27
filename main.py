from dotenv import load_dotenv
load_dotenv()  

from fastapi import FastAPI, Request, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from text_engine.pipeline import run
from adapters.docs_adapter import extract_text_from_bytes
from adapters.audio_adapter import process_audio
from adapters.audio_processor import apply_bleeps, convert_to_wav
import uvicorn
import io
import os
import re as _re
import json as _json
import tempfile
import shutil
from pathlib import Path

app = FastAPI(title="PrivacyShield")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

SUPPORTED_DOC_EXTS = {".txt", ".pdf", ".docx", ".csv", ".json", ".md", ".xls", ".xlsx"}


# ── TEXT TAB ──────────────────────────────────────────────
class TextInput(BaseModel):
    text: str

@app.post("/api/scan/text")
def scan_text(body: TextInput):
    return run(body.text, source="text")


# ── IMAGE TAB ─────────────────────────────────────────────
class ImageTextInput(BaseModel):
    extracted_text: str
    file_name: str = ""

_IMAGE_PATTERNS = [
    ("name",        _re.compile(r'\b[A-Z][a-z]+ [A-Z][a-z]+\b')),
    ("dob",         _re.compile(r'\b\d{1,2}[A-Za-z]{3}\d{4}\b|\b\d{2}/\d{2}/\d{4}\b')),
    ("expiry",      _re.compile(r'\b\d{1,2}[A-Za-z]{3}\d{4}\b')),
    ("card_number", _re.compile(r'\b\d{4,19}\b')),
    ("aadhaar",     _re.compile(r'\b\d{4}\s?\d{4}\s?\d{4}\b')),
    ("ssn",         _re.compile(r'\b\d{3}-\d{2}-\d{4}\b')),
    ("email",       _re.compile(r'\b[\w._%+\-]+@[\w.\-]+\.[a-zA-Z]{2,}\b')),
    ("phone",       _re.compile(r'\b(\(\d{3}\)\s?)?\d{3}[\s.\-]?\d{4}\b|\b\d{10}\b')),
]

@app.post("/api/scan/image")
def scan_image(body: ImageTextInput):
    result = run(body.extracted_text, source="image")
    has_values = any(
        d.get("value") or d.get("values")
        for d in result.get("detections", [])
    )
    if not has_values and result.get("risk_score", 0) > 0.3:
        extra_detections = []
        seen = set()
        for dtype, pattern in _IMAGE_PATTERNS:
            for m in pattern.finditer(body.extracted_text):
                val = m.group(0).strip()
                if val not in seen and len(val) > 1:
                    seen.add(val)
                    extra_detections.append({
                        "type":     dtype,
                        "value":    val,
                        "values":   [val],
                        "source":   "image_regex",
                        "severity": "HIGH"
                    })
        if extra_detections:
            result["detections"] = extra_detections
            result["action"] = "REDACT"
    print("=== FINAL IMAGE RESULT ===")
    print("detections:", result.get("detections"))
    print("==========================")
    return result


# ── DOCS TAB — Endpoint 1: pre-extracted text ─────────────
class DocsTextInput(BaseModel):
    extracted_text: str
    file_name: str = ""

@app.post("/api/scan/docs")
def scan_docs(body: DocsTextInput):
    return run(body.extracted_text, source="docs")


# ── DOCS TAB — Endpoint 2: raw file upload ────────────────
@app.post("/api/scan/docs-upload")
async def scan_docs_upload(file: UploadFile = File(...)):
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in SUPPORTED_DOC_EXTS:
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported type: {ext}. Supported: {', '.join(SUPPORTED_DOC_EXTS)}"
        )
    try:
        file_bytes = await file.read()
        extracted_text = extract_text_from_bytes(file_bytes, file.filename)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Text extraction failed: {str(e)}")

    if not extracted_text or not extracted_text.strip():
        raise HTTPException(status_code=422, detail="No text could be extracted from this file.")

    result = run(extracted_text, source="docs")
    result["file_name"]         = file.filename
    result["file_ext"]          = ext.lstrip(".")
    result["extracted_full"]    = extracted_text
    result["extracted_chars"]   = len(extracted_text)
    result["extracted_preview"] = extracted_text[:300] + ("…" if len(extracted_text) > 300 else "")
    return result


# ── DOCS TAB — Endpoint 3: format-preserving export ───────
@app.post("/api/export/redacted-doc")
async def export_redacted_doc(
    file: UploadFile = File(...),
    redacted_text: str = Form(...)
):
    ext        = os.path.splitext(file.filename or "")[1].lower()
    file_bytes = await file.read()
    try:
        if ext in (".txt", ".md", ".json"):
            output, mime, out_ext = redacted_text.encode("utf-8"), "text/plain", ext
        elif ext == ".csv":
            output, mime, out_ext = _export_csv(file_bytes, redacted_text)
        elif ext in (".xlsx", ".xls"):
            output, mime, out_ext = _export_xlsx(file_bytes, redacted_text)
        elif ext == ".docx":
            output, mime, out_ext = _export_docx(file_bytes, redacted_text)
        else:
            output, mime, out_ext = redacted_text.encode("utf-8"), "text/plain", ".txt"
    except Exception:
        output, mime, out_ext = redacted_text.encode("utf-8"), "text/plain", ".txt"

    base_name     = os.path.splitext(file.filename or "document")[0]
    download_name = f"{base_name}_redacted{out_ext}"
    return StreamingResponse(
        io.BytesIO(output),
        media_type=mime,
        headers={"Content-Disposition": f'attachment; filename="{download_name}"'}
    )


# ── Format export helpers ──────────────────────────────────
def _export_csv(original_bytes: bytes, redacted_text: str):
    import csv, io as _io
    orig_lines = original_bytes.decode("utf-8", errors="ignore").splitlines()
    red_lines  = redacted_text.splitlines()
    buf        = _io.StringIO()
    writer     = csv.writer(buf)
    for i, red_line in enumerate(red_lines):
        if i < len(orig_lines):
            orig_row = list(csv.reader([orig_lines[i]]))[0] if orig_lines[i] else []
            red_row  = list(csv.reader([red_line]))[0]      if red_line       else []
            if len(orig_row) > len(red_row):
                red_row += [""] * (len(orig_row) - len(red_row))
            writer.writerow(red_row[:len(orig_row)] if orig_row else red_row)
        else:
            writer.writerow([red_line])
    return buf.getvalue().encode("utf-8"), "text/csv", ".csv"

def _export_xlsx(original_bytes: bytes, redacted_text: str):
    import pandas as pd, io as _io
    try:
        orig_df   = pd.read_excel(_io.BytesIO(original_bytes))
        red_lines = redacted_text.splitlines()
        line_idx  = 0
        new_rows  = []
        for _, row in orig_df.iterrows():
            new_row = []
            for _ in row:
                new_row.append(red_lines[line_idx] if line_idx < len(red_lines) else "")
                line_idx += 1
            new_rows.append(new_row)
        new_df  = pd.DataFrame(new_rows, columns=orig_df.columns)
        out_buf = _io.BytesIO()
        new_df.to_excel(out_buf, index=False)
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        return out_buf.getvalue(), mime, ".xlsx"
    except Exception:
        return redacted_text.encode("utf-8"), "text/plain", ".txt"

def _export_docx(original_bytes: bytes, redacted_text: str):
    from docx import Document
    import io as _io
    try:
        orig_doc  = Document(_io.BytesIO(original_bytes))
        red_lines = redacted_text.splitlines()
        line_idx  = 0
        for para in orig_doc.paragraphs:
            if para.text.strip() and line_idx < len(red_lines):
                red_line = red_lines[line_idx]; line_idx += 1
                if para.runs:
                    para.runs[0].text = red_line
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.clear()
                    para.add_run(red_line)
        out_buf = _io.BytesIO()
        orig_doc.save(out_buf)
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return out_buf.getvalue(), mime, ".docx"
    except Exception:
        return redacted_text.encode("utf-8"), "text/plain", ".txt"


# ══════════════════════════════════════════════════════════
# AUDIO TAB
# ══════════════════════════════════════════════════════════

def _audio_to_wav(input_path: str, wav_path: str) -> str:
    """Convert any audio format to wav using imageio-ffmpeg."""
    import imageio_ffmpeg
    import subprocess
    ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
    subprocess.run(
        [ffmpeg_exe, "-y", "-i", input_path, wav_path],
        check=True, capture_output=True
    )
    return wav_path


@app.post("/api/audio/process")
async def audio_process(file: UploadFile = File(...)):
    """
    Step 1 of the new audio flow.
    Upload audio → Groq Whisper transcription → LLaMA PII detection
    → returns transcript + timestamped sensitive segments.
    Frontend uses this to show the toggle-chip region list.
    """
    suffix = Path(file.filename or "audio").suffix or ".wav"
    tmp_input = tempfile.mktemp(suffix=suffix)
    tmp_wav   = tempfile.mktemp(suffix=".wav")

    try:
        # Save upload
        with open(tmp_input, "wb") as f:
            shutil.copyfileobj(file.file, f)

        # Always convert to wav first (Groq needs PCM wav)
        if suffix.lower() != ".wav":
            print(f"[audio_process] Converting {suffix} → wav")
            _audio_to_wav(tmp_input, tmp_wav)
            process_path = tmp_wav
        else:
            process_path = tmp_input

        # Run full pipeline
        result = process_audio(process_path)

        return {
            "transcript":  result.get("transcript", ""),
            "action":      result.get("action", "ALLOW"),
            "segments":    result.get("segments", []),   # [{text,start,end,type}]
            "layers_used": result.get("layers_used", []),
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Audio processing failed: {str(e)}")
    finally:
        for p in [tmp_input, tmp_wav]:
            try: os.unlink(p)
            except: pass


@app.post("/api/audio/export")
async def audio_export(
    file:     UploadFile = File(...),
    segments: str        = Form(...),   # JSON: [{text,start,end,type}]
    format:   str        = Form("wav"), # "wav" or "mp3"
):
    """
    Step 2 of the new audio flow.
    Upload original audio + list of segments the user wants beeped
    → returns beeped audio file in requested format.
    Only segments the user left as 🔇 are sent here; revealed ones are excluded.
    """
    suffix    = Path(file.filename or "audio").suffix or ".wav"
    tmp_input = tempfile.mktemp(suffix=suffix)
    tmp_wav   = tempfile.mktemp(suffix=".wav")
    tmp_out   = tempfile.mktemp(suffix=".wav")
    tmp_mp3   = tempfile.mktemp(suffix=".mp3")

    try:
        # Save upload
        with open(tmp_input, "wb") as f:
            shutil.copyfileobj(file.file, f)

        # Convert to wav
        if suffix.lower() != ".wav":
            _audio_to_wav(tmp_input, tmp_wav)
            wav_path = tmp_wav
        else:
            wav_path = tmp_input

        # Parse segments
        segs = _json.loads(segments)
        print(f"[audio_export] Applying {len(segs)} bleep(s), format={format}")

        # Apply bleeps → wav
        apply_bleeps(wav_path, tmp_out, segs)

        # Convert to mp3 if requested
        if format == "mp3":
            _audio_to_wav(tmp_out, tmp_mp3)   # ffmpeg transcodes wav→mp3 too
            final_path = tmp_mp3
            mime       = "audio/mpeg"
            out_ext    = ".mp3"
        else:
            final_path = tmp_out
            mime       = "audio/wav"
            out_ext    = ".wav"

        with open(final_path, "rb") as f:
            audio_bytes = f.read()

        base          = os.path.splitext(file.filename or "audio")[0]
        download_name = f"{base}_redacted{out_ext}"

        return StreamingResponse(
            io.BytesIO(audio_bytes),
            media_type=mime,
            headers={"Content-Disposition": f'attachment; filename="{download_name}"'}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Audio export failed: {str(e)}")
    finally:
        for p in [tmp_input, tmp_wav, tmp_out, tmp_mp3]:
            try: os.unlink(p)
            except: pass


# ── Legacy audio endpoints (kept for backward compatibility) ──

@app.post("/api/scan/audio-file")
async def scan_audio_file(file: UploadFile = File(...)):
    """Legacy: transcribe + auto-beep all detected segments → return beeped wav."""
    tmp_dir = "temp"
    os.makedirs(tmp_dir, exist_ok=True)
    inp_path = os.path.join(tmp_dir, f"input_{file.filename}")
    wav_path = os.path.join(tmp_dir, "input_converted.wav")
    out_path = os.path.join(tmp_dir, "redacted_output.wav")

    try:
        with open(inp_path, "wb") as f:
            f.write(await file.read())
        import imageio_ffmpeg, subprocess
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        subprocess.run([ffmpeg_exe, "-y", "-i", inp_path, wav_path], check=True, capture_output=True)
        result = process_audio(wav_path)
        if result["action"] in ("REDACT", "BLOCK") and result["segments"]:
            apply_bleeps(wav_path, out_path, result["segments"])
            return FileResponse(out_path, media_type="audio/wav", filename=f"redacted_{file.filename}.wav")
        else:
            return {
                "action":     result["action"],
                "message":    "No sensitive content detected",
                "transcript": result["transcript"],
                "layers_used": result["layers_used"]
            }
    except Exception as e:
        raise HTTPException(500, str(e))
    finally:
        for path in [inp_path, wav_path]:
            if os.path.exists(path):
                os.remove(path)


@app.post("/api/debug-audio")
async def debug_audio(file: UploadFile = File(...)):
    """Debug: return raw transcript + timestamps without applying bleeps."""
    tmp_dir = "temp"
    os.makedirs(tmp_dir, exist_ok=True)
    inp_path = os.path.join(tmp_dir, f"input_{file.filename}")
    wav_path = os.path.join(tmp_dir, "debug_converted.wav")
    try:
        with open(inp_path, "wb") as f:
            f.write(await file.read())
        import imageio_ffmpeg, subprocess
        ffmpeg_exe = imageio_ffmpeg.get_ffmpeg_exe()
        subprocess.run([ffmpeg_exe, "-y", "-i", inp_path, wav_path], check=True, capture_output=True)
        return process_audio(wav_path)
    except Exception as e:
        raise HTTPException(500, str(e))
    finally:
        for path in [inp_path, wav_path]:
            if os.path.exists(path):
                os.remove(path)


# ── VIDEO TAB ─────────────────────────────────────────────
class VideoTextInput(BaseModel):
    extracted_text: str
    file_name: str = ""

@app.post("/api/scan/video")
def scan_video(body: VideoTextInput):
    return run(body.extracted_text, source="video")


# ── HEALTH CHECK ──────────────────────────────────────────
@app.get("/health")
def health():
    return {
        "status": "ok",
        "pipeline": ["regex", "distilbert", "gemini"],
        "endpoints": [
            "/api/scan/text",       "/api/scan/image",
            "/api/scan/docs",       "/api/scan/docs-upload",
            "/api/export/redacted-doc",
            "/api/audio/process",   "/api/audio/export",
            "/api/scan/audio-file", "/api/debug-audio",
            "/api/scan/video",
        ]
    }


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=False)
